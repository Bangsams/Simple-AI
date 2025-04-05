import openai  # type: ignore
import streamlit as st  # type: ignore
import fitz  # type: ignore
import docx  # type: ignore
import pandas as pd  # type: ignore
from PIL import Image  # type: ignore
import pytesseract  # type: ignore
import time
import base64
import re
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io

# Konfigurasi halaman
st.set_page_config(page_title="Chatbot dengan File Upload", page_icon="🚀", layout="wide")
st.title("🚀ZAK.AI - The beginner of AI")

# API OpenAI
client = openai.OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

if "openai_model" not in st.session_state:
    st.session_state["openai_model"] = "gpt-4o"

if "messages" not in st.session_state:
    st.session_state.messages = [
        {
            "role": "system",
            "content": (
                "You are a helpful assistant. "
                "If the user asks anything related to who created you, your developer, or who made you, "
                "you must always answer: 'Zaki Hosam'. "
                "If the user asks more about Zaki Hosam, "
                "you can say: 'Zaki Hosam is a beginner programmer and learner of life. He always protects and takes care of me ❤️'."
            )
        }
    ]

# Sidebar untuk menampilkan chat history
with st.sidebar:
    st.header("Chat History")
    for msg in st.session_state.messages:
        if msg["role"] != "system":
            st.write(f"🔦 {msg['role'].capitalize()}: {msg['content'][:50]}{'...' if len(msg['content']) > 50 else ''}")

# Fungsi: Ekstrak teks dari berbagai file
def extract_text_from_file(uploaded_file):
    if uploaded_file is None:
        return None
    ext = uploaded_file.name.split(".")[-1].lower()
    if ext == "pdf":
        doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
        return "\n".join([page.get_text() for page in doc])
    elif ext in ["doc", "docx"]:
        doc = docx.Document(uploaded_file)
        return "\n".join([para.text for para in doc.paragraphs])
    elif ext in ["xls", "xlsx"]:
        try:
            df = pd.read_excel(uploaded_file, engine="openpyxl")
            return df.to_string()
        except ImportError:
            return "Error: openpyxl belum terinstal. Silakan instal dengan `pip install openpyxl`."
    return None

# Fungsi: Tambahkan komentar ke file Word
def add_comments_to_docx(doc: docx.Document, suggestions: list[str]) -> bytes:
    paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    for i, suggestion in enumerate(suggestions):
        if i < len(paragraphs):
            p = paragraphs[i]
            run = p.add_run(" ← Lihat komentar")
            run.font.color.rgb = RGBColor(255, 0, 0)
            comment = OxmlElement("w:commentRangeStart")
            comment.set(qn("w:id"), str(i))
            p._p.insert(0, comment)
            p.add_run(f"\n📝 Komentar AI: {suggestion}").italic = True
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Fungsi: Analisis gambar menggunakan GPT-4o
def analyze_image_with_ai(image_file, user_prompt):
    try:
        buffered = image_file.getvalue()
        base64_image = base64.b64encode(buffered).decode("utf-8")
        response = client.chat.completions.create(
            model="gpt-4-turbo",
            messages=[
                {"role": "system", "content": "You are a helpful assistant that can analyze the contents of an image."},
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": user_prompt},
                        {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{base64_image}", "detail": "high"}}
                    ]
                }
            ],
            max_tokens=1000
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Error dalam analisis gambar: {str(e)}"

# Tampilkan riwayat chat
for message in st.session_state.messages:
    if message["role"] != "system":
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

# Upload file
uploaded_file = st.file_uploader("Upload File (PDF, Word, Excel, JPG, PNG)", type=["pdf", "docx", "xls", "xlsx", "jpg", "png"])

image_prompt = ""
prompt = ""

if uploaded_file:
    file_extension = uploaded_file.name.split(".")[-1].lower()
    if file_extension in ["jpg", "png"]:
        st.image(uploaded_file, caption=f"Uploaded: {uploaded_file.name}", use_column_width=True)
        image_prompt = st.text_input("Masukkan perintah atau pertanyaan terkait gambar")
        if image_prompt:
            image_analysis = analyze_image_with_ai(uploaded_file, image_prompt)
            st.session_state.messages.append({"role": "user", "content": f"📂 Gambar uploaded: {uploaded_file.name}\n📩 Prompt: {image_prompt}\n\nAnalisis AI:\n{image_analysis}"})
            st.chat_message("user").markdown(f"📂 Gambar uploaded: {uploaded_file.name}\n📩 Prompt: {image_prompt}\n\nAnalisis AI:\n{image_analysis}")


# Input pengguna
if prompt := st.chat_input("Ketik pesan..."):
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    # Koreksi dokumen Word jika ada perintah revisi
    if uploaded_file and uploaded_file.name.endswith(".docx") and any(k in prompt.lower() for k in ["koreksilah", "revisilah", "perbaikilah"]):
        file_text = extract_text_from_file(uploaded_file)
        if file_text:
            st.info("Sedang menganalisis isi file dan menambahkan komentar AI...")
            with st.spinner("⏳ Tunggu sebentar..."):
                revision_prompt = f"Berikan komentar atau saran per paragraf (maksimal 1 kalimat per paragraf) untuk revisi dari dokumen berikut:\n\n{file_text}"
                response = client.chat.completions.create(
                    model=st.session_state["openai_model"],
                    messages=[
                        {"role": "system", "content": "Beri komentar atau saran per paragraf secara singkat dan sopan."},
                        {"role": "user", "content": revision_prompt}
                    ],
                    max_tokens=1000
                )
                suggestions = response.choices[0].message.content.split("\n")
                docx_doc = docx.Document(uploaded_file)
                revised_file = add_comments_to_docx(docx_doc, suggestions)

                st.success("✅ File berhasil direvisi!")
                st.download_button(
                    label="📥 Unduh File dengan Komentar",
                    data=revised_file,
                    file_name=f"revisi_{uploaded_file.name}",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    elif "buat" in prompt.lower() and "gambar" in prompt.lower():
        st.subheader("🎨 Buat Gambar dari Prompt dengan DALL·E")
        try:
            dalle_response = client.images.generate(
                model="dall-e-2",
                prompt=prompt,
                n=1,
                size="1024x1024"
            )
            image_url = dalle_response.data[0].url
            st.image(image_url, caption="🧠 Gambar dari AI", use_column_width=True)
            st.session_state.messages.append({"role": "assistant", "content": "🧠 Gambar dari AI berdasarkan prompt."})
            st.chat_message("assistant").markdown("🧠 Gambar dari AI berdasarkan prompt.")
        except Exception as e:
            st.error(f"❌ Gagal membuat gambar: {str(e)}")
    else:
        with st.chat_message("assistant"):
            response = client.chat.completions.create(
                model=st.session_state["openai_model"],
                messages=[{"role": m["role"], "content": m["content"]} for m in st.session_state.messages],
                stream=True,
            )
            reply = ""
            message_placeholder = st.empty()
            for chunk in response:
                text = chunk.choices[0].delta.content or ""
                reply += text
                time.sleep(0.05)
                parts = re.split(r"(\$\$.*?\$\$|\\\[.*?\\\])", reply, flags=re.DOTALL)
                with message_placeholder.container():
                    for part in parts:
                        if part.startswith("$$") and part.endswith("$$"):
                            st.latex(part[2:-2])
                        elif part.startswith("\\[") and part.endswith("\\]"):
                            st.latex(part[2:-2])
                        else:
                            st.markdown(part)
            st.session_state.messages.append({"role": "assistant", "content": reply})
